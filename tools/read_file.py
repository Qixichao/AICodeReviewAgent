import PyPDF2
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders.word_document import UnstructuredWordDocumentLoader
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
from docx import Document


def read_from_file(filepath: str):
    ext = filepath.split('.')[-1]
    if ext == "pdf":
        return PyPDFLoader(filepath)
    elif ext == "docx" or ext == "doc":
        return UnstructuredWordDocumentLoader(filepath)
    else:
        raise NotImplementedError(f"File extension {ext} not supported.")

def read_from_file_raw(filepath: str):
    ext = filepath.split('.')[-1]
    if ext == "pdf":
        return pdf_to_text(filepath)
    elif ext == "docx" or ext == "doc":
        return read_item_from_word_element(filepath)
    else:
        raise NotImplementedError(f"File extension {ext} not supported.")

# def read_item_from_pdf(filepath: str):
#     with open(filepath, 'rb') as file:
#         reader = PyPDF2.PdfReader(file)
#         num_pages = len(reader.pages)
#         for i in range(num_pages):
#             page = reader.pages[i]
#             # 处理每一页的内容
#             text = page.extract_text()
#             print(text)



def pdf_to_text(path):
    manager = PDFResourceManager()
    output = StringIO()
    laparams = LAParams()
    converter = TextConverter(manager, output, laparams=laparams)
    interpreter = PDFPageInterpreter(manager, converter)

    with open(path, 'rb') as file:
        for page in PDFPage.get_pages(file, caching=True, check_extractable=True):
            interpreter.process_page(page)

        filecontent = output.getvalue()

    converter.close()
    output.close()
    return filecontent


# def read_item_from_word(path):
#     # 打开Word文档
#     document = Document(path)
#     filecontent = ""
#     tablecontent = ""
#     # 读取每一段落的内容
#     for paragraph in document.paragraphs:
#         filecontent = filecontent + paragraph.text
#
#     # 读取每一表格的内容
#     for table in document.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 print(cell.text)
#     return filecontent

def read_item_from_word_element(path):
    # 打开Word文档
    document = Document(path)

    # 读取所有内容
    filecontent = ''
    for element in document.element.body:
        if element.tag.endswith('p'):
            filecontent = filecontent + element.text + "\n"
        elif element.tag.endswith('tbl'):
            for tbl_element in element.iter():
                if tbl_element.tag.endswith('t'):
                    filecontent += tbl_element.text

    return filecontent


#print(read_from_file_raw('../docs/代码编写规范.pdf'))
#print(read_from_file_raw('../docs/代码编写规范.docx'))
#pdf_to_text('../docs/代码编写规范.pdf')